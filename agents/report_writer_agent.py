"""
Report Writer Agent
Generates comprehensive Zero Trust assessment reports in DOCX format
"""

from typing import Dict, Any, List
from datetime import datetime
from .base_agent import BaseAgent

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass  # Will be installed in requirements


class ReportWriterAgent(BaseAgent):
    """
    Generates professional assessment reports.
    
    Uses two-model approach:
    - Claude Haiku 4.5: Executive summary (fast, efficient)
    - Claude Sonnet 4: Technical findings (detailed, accurate)
    """
    
    def __init__(self, **kwargs):
        super().__init__(
            name="ReportWriterAgent",
            model="claude-sonnet-4-20250514",  # Default to Sonnet
            **kwargs
        )
        
        # Create Haiku client for summaries
        from langchain_anthropic import ChatAnthropic
        import os
        
        self.haiku_llm = ChatAnthropic(
            model="claude-haiku-4-5-20241001",
            temperature=0.0,
            anthropic_api_key=os.getenv("ANTHROPIC_API_KEY"),
            max_tokens=2048
        )
    
    async def process(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate comprehensive report.
        
        Expected state keys:
            - zt_scores: Zero Trust pillar scores
            - zt_gaps: List of gaps
            - zt_strengths: List of strengths
            - zt_recommendations: List of recommendations
            - compliance_matrix: Compliance framework results
            - overall_maturity_level: Maturity level name
            
        Returns updated state with:
            - report_path: Path to generated DOCX file
            - executive_summary: Text of executive summary
            - report_generated: True
        """
        self.logger.info("Generating assessment report")
        
        # Validate input
        required = ['zt_scores', 'overall_maturity_level', 'compliance_matrix']
        self.validate_state(state, required)
        
        # Generate executive summary (Haiku - fast)
        exec_summary = await self._generate_executive_summary(state)
        
        # Create DOCX document
        doc = Document()
        
        # Add report sections
        self._add_cover_page(doc, state)
        self._add_executive_summary(doc, exec_summary)
        self._add_maturity_scores(doc, state)
        self._add_findings(doc, state)
        self._add_compliance_section(doc, state)
        self._add_recommendations(doc, state)
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_filename = f"ZeroTrust_Assessment_{timestamp}.docx"
        report_path = f"/home/claude/{report_filename}"
        
        doc.save(report_path)
        
        self.logger.info(f"Report generated: {report_path}")
        
        # Update state
        updates = {
            'report_path': report_path,
            'report_filename': report_filename,
            'executive_summary': exec_summary,
            'report_generated': True
        }
        
        return self.update_state(state, updates)
    
    async def _generate_executive_summary(
        self,
        state: Dict[str, Any]
    ) -> str:
        """Generate executive summary using Haiku."""
        
        system_prompt = """You are writing an executive summary for a Zero Trust security assessment.

Write a concise 3-4 paragraph summary that includes:
1. Overall assessment (maturity level and score)
2. Key strengths (2-3 highlights)
3. Critical gaps (2-3 priorities)
4. Strategic recommendations (high-level)

Target audience: C-level executives and security leadership.
Tone: Professional, clear, action-oriented.
Length: 300-400 words."""
        
        # Build context
        maturity = state.get('overall_maturity_level', 'Unknown')
        score = state.get('overall_maturity_score', 0)
        compliance_pct = state.get('compliance_percentage', 0)
        
        strengths = state.get('zt_strengths', [])[:5]
        gaps = state.get('zt_gaps', [])[:5]
        
        context = f"""
OVERALL MATURITY: {maturity} ({score}/5.0)
COMPLIANCE: {compliance_pct}%

TOP STRENGTHS:
{chr(10).join(['- ' + s for s in strengths])}

CRITICAL GAPS:
{chr(10).join(['- ' + g for g in gaps])}

Write the executive summary."""
        
        # Use Haiku for speed
        from langchain_core.messages import HumanMessage, SystemMessage
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=context)
        ]
        
        response = await self.haiku_llm.ainvoke(messages)
        
        return response.content
    
    def _add_cover_page(
        self,
        doc: Document,
        state: Dict[str, Any]
    ) -> None:
        """Add cover page."""
        
        # Title
        title = doc.add_heading('Zero Trust Architecture Assessment', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        doc.add_paragraph()
        subtitle = doc.add_paragraph('Comprehensive Security Posture Analysis')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Metadata
        doc.add_paragraph()
        doc.add_paragraph()
        
        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_str = datetime.now().strftime("%B %d, %Y")
        maturity = state.get('overall_maturity_level', 'Unknown')
        
        metadata.add_run(f"Assessment Date: {date_str}\n")
        metadata.add_run(f"Overall Maturity: {maturity}\n")
        metadata.add_run("Generated by VaultZero v2.0")
        
        # Page break
        doc.add_page_break()
    
    def _add_executive_summary(
        self,
        doc: Document,
        summary: str
    ) -> None:
        """Add executive summary section."""
        
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(summary)
        doc.add_page_break()
    
    def _add_maturity_scores(
        self,
        doc: Document,
        state: Dict[str, Any]
    ) -> None:
        """Add maturity scores table."""
        
        doc.add_heading('Zero Trust Maturity Assessment', level=1)
        
        zt_scores = state.get('zt_scores', {})
        
        # Add summary
        overall_score = state.get('overall_maturity_score', 0)
        maturity_level = state.get('overall_maturity_level', 'Unknown')
        
        summary = doc.add_paragraph()
        summary.add_run(f"Overall Maturity: {overall_score}/5.0 ({maturity_level})")
        summary.runs[0].font.bold = True
        summary.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Zero Trust Pillar'
        header_cells[1].text = 'Score'
        header_cells[2].text = 'Maturity Level'
        
        # Make header bold
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add scores
        for pillar, score in zt_scores.items():
            row_cells = table.add_row().cells
            row_cells[0].text = pillar
            row_cells[1].text = f"{score}/5"
            row_cells[2].text = self._score_to_level(score)
        
        doc.add_page_break()
    
    def _add_findings(
        self,
        doc: Document,
        state: Dict[str, Any]
    ) -> None:
        """Add findings section."""
        
        doc.add_heading('Key Findings', level=1)
        
        # Strengths
        doc.add_heading('Strengths', level=2)
        strengths = state.get('zt_strengths', [])
        
        if strengths:
            for strength in strengths[:10]:  # Top 10
                doc.add_paragraph(strength, style='List Bullet')
        else:
            doc.add_paragraph("No significant strengths identified.")
        
        doc.add_paragraph()
        
        # Gaps
        doc.add_heading('Gaps & Weaknesses', level=2)
        gaps = state.get('zt_gaps', [])
        
        if gaps:
            for gap in gaps[:15]:  # Top 15
                doc.add_paragraph(gap, style='List Bullet')
        else:
            doc.add_paragraph("No critical gaps identified.")
        
        doc.add_page_break()
    
    def _add_compliance_section(
        self,
        doc: Document,
        state: Dict[str, Any]
    ) -> None:
        """Add compliance mapping section."""
        
        doc.add_heading('Compliance Framework Mapping', level=1)
        
        compliance_matrix = state.get('compliance_matrix', {})
        overall_compliance = state.get('compliance_percentage', 0)
        
        # Overall compliance
        summary = doc.add_paragraph()
        summary.add_run(f"Overall Compliance: {overall_compliance}%")
        summary.runs[0].font.bold = True
        summary.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Framework details
        for framework_id, result in compliance_matrix.items():
            doc.add_heading(framework_id, level=2)
            
            score = result.get('compliance_score', 0)
            total = result.get('total_controls', 0)
            met = result.get('controls_met_count', 0)
            
            doc.add_paragraph(f"Compliance Score: {score}% ({met}/{total} controls met)")
            
            # Controls met
            controls_met = result.get('controls_met', [])
            if controls_met:
                doc.add_paragraph("Controls Satisfied:", style='List Bullet')
                for control in controls_met[:5]:  # Top 5
                    doc.add_paragraph(control.replace(f"[{framework_id}] ", ""), 
                                    style='List Bullet 2')
            
            # Gaps
            framework_gaps = result.get('gaps', [])
            if framework_gaps:
                doc.add_paragraph("Gaps:", style='List Bullet')
                for gap in framework_gaps[:5]:  # Top 5
                    doc.add_paragraph(gap.replace(f"[{framework_id}] ", ""), 
                                    style='List Bullet 2')
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_recommendations(
        self,
        doc: Document,
        state: Dict[str, Any]
    ) -> None:
        """Add recommendations section."""
        
        doc.add_heading('Recommendations', level=1)
        
        recommendations = state.get('zt_recommendations', [])
        
        if recommendations:
            # Prioritize recommendations
            doc.add_paragraph("The following recommendations are prioritized "
                            "based on impact and feasibility:")
            doc.add_paragraph()
            
            for i, rec in enumerate(recommendations[:20], 1):  # Top 20
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
        else:
            doc.add_paragraph("No specific recommendations generated.")
    
    def _score_to_level(self, score: float) -> str:
        """Convert score to maturity level."""
        if score < 1.0:
            return "Initial"
        elif score < 2.0:
            return "Developing"
        elif score < 3.0:
            return "Defined"
        elif score < 4.0:
            return "Managed"
        else:
            return "Optimized"
